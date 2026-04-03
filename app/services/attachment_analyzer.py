"""
attachment_analyzer.py — Static + Dynamic phishing analysis for file attachments.
Phase 8: Static inspection (extension, filename, URL extraction, text NLP).
Phase 10: Dynamic VirusTotal analysis — hash lookup, file upload, polling, URL scanning.
"""
import asyncio
import hashlib
import io
import os
import time

import requests

from app.services.feature_extractor import URLFeatureExtractor
from app.services.model_runner import ModelRunner
from app.services.text_analyzer import TextAnalyzer

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_URLS_TO_SCAN = 10
HIGH_RISK_URL_THRESHOLD = 2  # 2+ HIGH_RISK URLs -> PHISHING override
VT_UPLOAD_MAX_BYTES = 32 * 1024 * 1024  # 32 MB VT upload limit

DANGEROUS_EXTENSIONS = {".exe", ".js", ".vbs", ".bat", ".ps1", ".scr", ".cmd"}
ACCEPTED_EXTENSIONS = DANGEROUS_EXTENSIONS | {".pdf", ".docx"}
SUSPICIOUS_FILENAME_KEYWORDS = ["invoice", "payment", "urgent", "verify", "receipt"]


class AttachmentAnalyzer:
    def __init__(self, url_extractor: URLFeatureExtractor, model_runner: ModelRunner, text_analyzer: TextAnalyzer = None):
        self.url_extractor = url_extractor
        self.model_runner = model_runner
        self.text_analyzer = text_analyzer

    def _url_to_score(self, risk: str, conf: float) -> float:
        if risk == "HIGH_RISK":
            return min(100.0, max(75.0, conf * 100))
        elif risk == "SUSPICIOUS":
            return 45.0 + (conf * 29.0)
        else:
            return max(0.0, min(40.0, (1.0 - conf) * 40))

    def _extract_urls_from_pdf(self, file_bytes: bytes) -> list[str]:
        """Extract all embedded URLs from a PDF using PyMuPDF (fitz)."""
        urls = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                for link in page.get_links():
                    uri = link.get("uri", "")
                    if uri and uri.startswith("http"):
                        urls.append(uri)
            doc.close()
        except ImportError:
            pass  # PyMuPDF not installed
        except Exception:
            pass  # Corrupt or unreadable PDF
        return list(dict.fromkeys(urls))  # deduplicate preserving order

    def _extract_urls_from_docx(self, file_bytes: bytes) -> list[str]:
        """Extract hyperlinks from a DOCX file using python-docx."""
        urls = []
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel.target_ref
                    if url and url.startswith("http"):
                        urls.append(url)
        except ImportError:
            pass  # python-docx not installed
        except Exception:
            pass  # Corrupt or unreadable DOCX
        return list(dict.fromkeys(urls))  # deduplicate

    def _extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract up to 5 pages / 10,000 characters from a PDF."""
        text = ""
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc[:5]:
                text += page.get_text() + "\n"
            doc.close()
        except Exception:
            pass
        return text[:10000]

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract up to 10,000 characters from a DOCX."""
        text = ""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
                if len(text) > 10000:
                    break
        except Exception:
            pass
        return text[:10000]

    # =========================================================================
    # Phase 10: VirusTotal Dynamic Analysis
    # =========================================================================

    async def _run_virustotal_analysis(self, file_bytes: bytes, filename: str) -> tuple[int, list[str], dict]:
        """
        Dynamic analysis via VirusTotal API v3.

        Flow:
          1. SHA-256 hash lookup  → GET /api/v3/files/{hash}
          2. If 404 + upload enabled + file ≤32MB → POST /api/v3/files → poll analysis
          3. Return (vt_score, findings, vt_meta)

        Granular scoring:
          malicious >= 3  → 100
          malicious >= 1  → 85
          suspicious >= 3 → 70
          suspicious >= 1 → 50
          else            → 0
        """
        vt_key = os.getenv("VIRUSTOTAL_API_KEY")
        if not vt_key:
            return 0, ["VirusTotal API key not configured — dynamic scan skipped"], {
                "active": False, "stats": None, "link": None
            }

        headers = {"x-apikey": vt_key, "accept": "application/json"}
        file_hash = hashlib.sha256(file_bytes).hexdigest()
        findings: list[str] = []
        vt_score = 0
        vt_stats = None
        vt_link = None

        def _score_from_stats(stats: dict) -> int:
            malicious = stats.get("malicious", 0)
            suspicious = stats.get("suspicious", 0)
            if malicious >= 3:    return 100
            elif malicious >= 1:  return 85
            elif suspicious >= 3: return 70
            elif suspicious >= 1: return 50
            else:                 return 0

        def _poll_analysis(analysis_id: str) -> dict | None:
            """Poll GET /api/v3/analyses/{id} every 15s, max 12 attempts (180s timeout)."""
            poll_url = f"https://www.virustotal.com/api/v3/analyses/{analysis_id}"
            for attempt in range(12):
                time.sleep(15)
                try:
                    resp = requests.get(poll_url, headers=headers, timeout=30)
                    if resp.status_code == 200:
                        data = resp.json().get("data", {})
                        status = data.get("attributes", {}).get("status", "")
                        if status == "completed":
                            return data.get("attributes", {}).get("stats", {})
                except Exception:
                    pass
            return None  # timed out

        try:
            # --- Step 1: Hash lookup ---
            resp = await asyncio.to_thread(
                lambda: requests.get(
                    f"https://www.virustotal.com/api/v3/files/{file_hash}",
                    headers=headers,
                    timeout=30
                )
            )

            if resp.status_code == 200:
                data = resp.json()
                vt_stats = data.get("data", {}).get("attributes", {}).get("last_analysis_stats", {})
                vt_score = _score_from_stats(vt_stats)
                vt_link = f"https://www.virustotal.com/gui/file/{file_hash}"
                malicious = vt_stats.get("malicious", 0)
                suspicious = vt_stats.get("suspicious", 0)
                if malicious > 0 or suspicious > 0:
                    findings.append(
                        f"VirusTotal [HASH HIT]: {malicious} malicious, {suspicious} suspicious across 70+ AV engines."
                    )
                else:
                    findings.append("VirusTotal [HASH HIT]: 0 engines flagged this file (Clean).")

            elif resp.status_code == 404:
                # --- Step 2: Upload if allowed ---
                vt_upload_enabled = os.getenv("VT_UPLOAD_ENABLED", "true").lower() != "false"
                if vt_upload_enabled and len(file_bytes) <= VT_UPLOAD_MAX_BYTES:
                    findings.append("VirusTotal: Hash unknown — uploading file for dynamic analysis...")
                    upload_resp = await asyncio.to_thread(
                        lambda: requests.post(
                            "https://www.virustotal.com/api/v3/files",
                            headers={"x-apikey": vt_key},
                            files={"file": (filename, file_bytes)},
                            timeout=60
                        )
                    )
                    if upload_resp.status_code == 200:
                        analysis_id = upload_resp.json().get("data", {}).get("id", "")
                        if analysis_id:
                            findings.append(
                                f"VirusTotal: Upload queued (analysis_id={analysis_id[:12]}...). "
                                "Polling for results (max 180s)..."
                            )
                            stats = await asyncio.to_thread(_poll_analysis, analysis_id)
                            if stats is not None:
                                vt_stats = stats
                                vt_score = _score_from_stats(stats)
                                vt_link = f"https://www.virustotal.com/gui/file/{file_hash}"
                                malicious = stats.get("malicious", 0)
                                suspicious = stats.get("suspicious", 0)
                                findings.append(
                                    f"VirusTotal [UPLOAD SCAN]: {malicious} malicious, {suspicious} suspicious."
                                )
                            else:
                                findings.append(
                                    "VirusTotal: Scan timed out after 180s — dynamic result unavailable."
                                )
                        else:
                            findings.append("VirusTotal: Upload succeeded but no analysis_id returned.")
                    elif upload_resp.status_code == 429:
                        findings.append("VirusTotal rate limit reached (4 req/min) — upload skipped.")
                    else:
                        findings.append(f"VirusTotal upload failed: HTTP {upload_resp.status_code}.")
                elif len(file_bytes) > VT_UPLOAD_MAX_BYTES:
                    findings.append("VirusTotal: File too large for upload (>32MB) — dynamic scan skipped.")
                else:
                    findings.append("VirusTotal: Hash unknown, upload disabled (VT_UPLOAD_ENABLED=false).")

            elif resp.status_code == 429:
                findings.append("VirusTotal rate limit reached (4 req/min) — dynamic scan skipped.")

            elif resp.status_code == 401:
                findings.append("VirusTotal API key invalid or expired — dynamic scan skipped.")

            else:
                findings.append(f"VirusTotal API error: HTTP {resp.status_code}.")

        except Exception as e:
            findings.append(f"VirusTotal error: {str(e)[:120]}")

        return vt_score, findings, {"active": True, "stats": vt_stats, "link": vt_link}

    async def _scan_url_via_vt(self, target_url: str) -> dict:
        """
        Submit a URL to VirusTotal for scanning.
        POST /api/v3/urls → poll GET /api/v3/analyses/{id} (6 × 10s = 60s max).
        Returns dict with vt_checked, vt_engines_malicious, vt_engines_suspicious.
        """
        vt_key = os.getenv("VIRUSTOTAL_API_KEY")
        if not vt_key:
            return {"vt_checked": False}

        headers_form = {"x-apikey": vt_key, "content-type": "application/x-www-form-urlencoded"}
        headers_get = {"x-apikey": vt_key, "accept": "application/json"}

        try:
            resp = await asyncio.to_thread(
                lambda: requests.post(
                    "https://www.virustotal.com/api/v3/urls",
                    headers=headers_form,
                    data={"url": target_url},
                    timeout=30
                )
            )
            if resp.status_code == 429:
                return {"vt_checked": False, "vt_rate_limited": True}
            if resp.status_code != 200:
                return {"vt_checked": False}

            analysis_id = resp.json().get("data", {}).get("id", "")
            if not analysis_id:
                return {"vt_checked": False}

            # Poll for completion — 6 attempts × 10s = 60s timeout
            poll_url = f"https://www.virustotal.com/api/v3/analyses/{analysis_id}"
            _hdrs = headers_get  # capture for lambda
            for _ in range(6):
                time.sleep(10)
                try:
                    poll_resp = await asyncio.to_thread(
                        lambda: requests.get(poll_url, headers=_hdrs, timeout=30)
                    )
                    if poll_resp.status_code == 200:
                        data = poll_resp.json().get("data", {})
                        status = data.get("attributes", {}).get("status", "")
                        if status == "completed":
                            stats = data.get("attributes", {}).get("stats", {})
                            return {
                                "vt_checked": True,
                                "vt_engines_malicious": stats.get("malicious", 0),
                                "vt_engines_suspicious": stats.get("suspicious", 0),
                            }
                except Exception:
                    pass

            return {
                "vt_checked": True,
                "vt_engines_malicious": 0,
                "vt_engines_suspicious": 0,
                "vt_timed_out": True,
            }

        except Exception:
            return {"vt_checked": False}

    # =========================================================================
    # Core analysis — now async (Phase 10)
    # =========================================================================

    async def analyze(self, filename: str, file_bytes: bytes) -> dict:
        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            raise ValueError("File too large. Maximum size is 10MB.")

        _, ext = os.path.splitext(filename.lower())
        if ext not in ACCEPTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Accepted types: .pdf, .docx, .exe, .js, .vbs, .bat, .ps1, .scr, .cmd"
            )

        findings = []
        ext_score = 0
        if ext in DANGEROUS_EXTENSIONS:
            ext_score = 100
            findings.append(f"Dangerous file extension detected: '{ext}' — executable or script type")

        filename_score = 0
        name_lower = os.path.splitext(filename.lower())[0]
        for keyword in SUSPICIOUS_FILENAME_KEYWORDS:
            if keyword in name_lower:
                filename_score = 60
                findings.append(f"Suspicious filename keyword detected: '{keyword}'")
                break  # one flag is enough

        embedded_urls = []
        extracted_text = ""
        if ext == ".pdf":
            embedded_urls = self._extract_urls_from_pdf(file_bytes)
            extracted_text = self._extract_text_from_pdf(file_bytes)
        elif ext == ".docx":
            embedded_urls = self._extract_urls_from_docx(file_bytes)
            extracted_text = self._extract_text_from_docx(file_bytes)

        text_score = 0.0
        if extracted_text and self.text_analyzer:
            text_res = self.text_analyzer.analyze(extracted_text)
            text_score = text_res.get("risk_score", 0)
            if text_res.get("verdict") == "PHISHING":
                findings.append("Document content flagged as PHISHING by roberta-spam semantics.")
            elif text_res.get("verdict") == "SUSPICIOUS":
                findings.append("Document content flagged as SUSPICIOUS by NLP model.")

        # Cap to MAX_URLS_TO_SCAN
        embedded_urls = embedded_urls[:MAX_URLS_TO_SCAN]
        url_scan_results = []
        high_risk_count = 0

        for url in embedded_urls:
            try:
                features = self.url_extractor.extract(url)
                features["_raw_url"] = url  # enable FIX-01 keyword detection
                result = self.model_runner.predict(features)
                risk = result.get("risk_flag", "SAFE")
                if risk == "HIGH_RISK":
                    high_risk_count += 1

                # Phase 10: VT URL enrichment (sequential — rate limit aware)
                vt_url_info = await self._scan_url_via_vt(url)
                if vt_url_info.get("vt_rate_limited"):
                    findings.append("VirusTotal URL scan rate limited — remaining URLs skipped.")
                    embedded_urls = []  # stop further VT URL calls
                    vt_url_info = {"vt_checked": False}

                url_scan_results.append({
                    "url": url,
                    "risk_flag": risk,
                    "confidence": result.get("confidence"),
                    **vt_url_info,
                })
                findings.append(
                    f"Embedded URL {'[HIGH RISK]' if risk == 'HIGH_RISK' else '[' + risk + ']'}: {url}"
                )
            except Exception:
                pass  # skip unscanneable URLs silently

        if url_scan_results:
            url_score = sum(
                self._url_to_score(r["risk_flag"], r.get("confidence", 0.5))
                for r in url_scan_results
            ) / len(url_scan_results)
        else:
            url_score = 0.0

        # === Dynamic Behavior Check via VirusTotal (Phase 10) ===
        vt_score, vt_findings, vt_meta = await self._run_virustotal_analysis(file_bytes, filename)
        findings.extend(vt_findings)

        # === Score Calculation ===
        # Weights: ext(20%) + filename(10%) + url(25%) + text(20%) + vt(25%)
        attachment_score = round(
            (ext_score * 0.20)
            + (filename_score * 0.10)
            + (url_score * 0.25)
            + (text_score * 0.20)
            + (vt_score * 0.25)
        )

        # Hard overrides
        if ext_score == 100:
            attachment_score = max(attachment_score, 80)
        if vt_score >= 75:
            attachment_score = max(attachment_score, vt_score)

        attachment_score = max(0, min(100, attachment_score))

        if vt_score == 100:
            verdict = "PHISHING"
            findings.insert(0, "CRITICAL: VirusTotal dynamic sandbox confirms definitive malware/phishing.")
        elif high_risk_count >= HIGH_RISK_URL_THRESHOLD:
            verdict = "PHISHING"  # hard override
            findings.insert(0, f"CRITICAL: {high_risk_count} embedded URLs are HIGH_RISK — multi-URL phishing chain detected")
        elif attachment_score > 70:
            verdict = "PHISHING"
        elif attachment_score >= 40:
            verdict = "SUSPICIOUS"
        else:
            verdict = "SAFE"

        return {
            "filename": filename,
            "attachment_score": attachment_score,
            "verdict": verdict,
            "findings": findings,
            "breakdown": {
                "extension_score": ext_score,
                "filename_score": filename_score,
                "url_score": round(url_score, 1),
                "text_score": round(text_score, 1),
                "virustotal_cloud_score": vt_score,
                "high_risk_url_count": high_risk_count,
                "embedded_urls_scanned": len(url_scan_results),
            },
            "url_scan_results": url_scan_results,
            "virustotal": vt_meta,
        }
